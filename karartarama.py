import os, sys, csv, math, datetime
# import PyPDF2
import tkinter
from tkinter import filedialog
from os import path
import docx
import PyPDF2
from tika import parser
import re
from docx import Document
from docx.shared import Pt



root = tkinter.Tk()
root.withdraw() #use to hide tkinter window

def search_for_file_path ():
    currdir = os.getcwd()
    tempdir = filedialog.askdirectory(parent=root, initialdir=currdir, title='Lutfen hedef klasoru seciniz.')
    if len(tempdir) > 0:
        print ("You chose: %s" % tempdir)
    return tempdir

def hasNumbers(inputString):
    return any(char.isdigit() for char in inputString)

def makeDocument(filename,outputpath, daireno, esasno, kararno, tarih, metin):
    filename1=os.path.splitext(filename)[0]
    output_path1=output_path+ "/sablon-" + filename1 + ".docx"
    outputdocument=docx.Document( )


    style = outputdocument.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size=Pt(12)
    
    paragraph0=outputdocument.add_paragraph("")

    run0 = paragraph0.add_run("T.C. YARGITAY  \n" + daireno + ' Hukuk Dairesi')
    run0.bold = True
    
    table= outputdocument.add_table(3, 3)
  

    row0 = table.rows[0]
    row0.cells[0].text = 'Esas No. '
    row0.cells[1].text = esasno
    row0.cells[2].text='İlgili Kanun/Madde:'

    row1=table.rows[1]
    row1.cells[0].text = 'Karar No. '
    row1.cells[1].text = kararno


    row2=table.rows[2]
    row2.cells[0].text = 'Tarihi:'
    row2.cells[1].text = tarih

    

    
    style = outputdocument.styles['No Spacing']
    font = style.font
    font.name = 'Verdana'
    font.size=Pt(8)
    

    paragraph1 = outputdocument.add_paragraph(metin)
    format1=paragraph1.paragraph_format
    format1.alignment=docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    paragraph1.style = outputdocument.styles['No Spacing']
    
    
    
    outputdocument.save(output_path1)


pdfFiles = []




file_path_variable = search_for_file_path()
print ("\nfile_path_variable = ", file_path_variable)
# define the name of the directory to be created
output_path = file_path_variable+ "/sablon_kararlar"
if not path.exists(output_path):
    try:
        os.mkdir(output_path)
    except OSError:
        print ("Creation of the directory %s failed" % output_path)
    else:
        print ("Successfully created the directory %s " % output_path)




for filename in os.listdir(file_path_variable):
    if filename.endswith('.pdf'):
        pdfFiles.append(filename)
        
# pdfFiles.sort(key=str.lower)
for filename in pdfFiles:
    # filename1=os.path.splitext(filename)[0]
    # output_path1=output_path+"/sablon-"+filename1+".docx"
    # outputdocument=docx.Document()
    # outputdocument.add_heading("T.C. YARGITAY", 1)
    # outputdocument.save(output_path1)
    raw=parser.from_file(file_path_variable+"/"+filename)
    string1=raw['content']
    a_list=str.split(string1)
    esasno=a_list[a_list.index('E.')-1]
    kararno=a_list[a_list.index('K.')-1]
    hukukdairesino=a_list[a_list.index('Dairesi')-2]
    index1=a_list.index('verildi.')
    
    karartarihi=''
    for i in range(len(a_list)-1, 1, -1):
        if a_list[i]=='verildi.':
            index1=i
            for k in range(i,1,-1):
                if hasNumbers(a_list[k]):
                    karartarihi= a_list[k]
                    break

    
    
    baslangic = string1.index('"İçtihat Metni"')
    son=string1.index('karar verildi.')+14


    kararmetni = string1[baslangic:son]

    makeDocument(file_path_variable,filename,output_path, hukukdairesino, esasno, kararno, karartarihi, kararmetni)
    #this code has more cool to stuff on the way


                    
    # print(karartarihi)

    # //print(string1.lower().index('karar verildi'))
    # re.search()
    # m = re.search("^([1-9] |1[0-9]| 2[0-9]|3[0-1])(.-)([1-9] |1[0-2])(./-|)20[0-9][0-9]$", string1)
    # prin"t(m)
    # date=""
    # for date in a_list:
    #^((0?[1-9]|[1-2][0-9]|3[0-1])(/|-|.)(0?[1-9]|1[0-2])(/|-|.)([1-2][0-9][0-9][0-9]))?$
    # x=re.search(r'((0?[1-9]|[1-2][0-9]|3[0-1])/(0?[1-9]|1[0-2])/([1-2][0-9][0-9][0-9]))', string1)
    # print(x[3]) 

        

    # x=re.search("karar",string1)
    
   
    # objectfile=open(file_path_variable+"/"+filename, 'rb')
    # pdfReader=PyPDF2.PdfFileReader(objectfile)
  

   

    
   

